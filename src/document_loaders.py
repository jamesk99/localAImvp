"""
Document loaders for various file formats.
Each loader extracts text content from a specific file type.
"""
import os
from pathlib import Path
from typing import Optional, Dict, Any
import PyPDF2


class DocumentLoader:
    """Base class for document loaders."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        """Load and extract text from a document. Returns None if failed."""
        raise NotImplementedError


class TextLoader(DocumentLoader):
    """Loader for plain text files (.txt, .md)."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"   ❌ Error loading text file: {e}")
            return None


class PDFLoader(DocumentLoader):
    """Loader for PDF files."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"   ❌ Error loading PDF: {e}")
            return None


class DocxLoader(DocumentLoader):
    """Loader for Microsoft Word documents (.docx)."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            from docx import Document
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)
            
            return "\n".join(text)
        except ImportError:
            print(f"   ⚠️  python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            print(f"   ❌ Error loading DOCX: {e}")
            return None


class CSVLoader(DocumentLoader):
    """Loader for CSV files."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            import csv
            text = []
            
            with open(file_path, 'r', encoding='utf-8') as f:
                # Try to detect delimiter
                sample = f.read(1024)
                f.seek(0)
                sniffer = csv.Sniffer()
                try:
                    delimiter = sniffer.sniff(sample).delimiter
                except:
                    delimiter = ','
                
                reader = csv.reader(f, delimiter=delimiter)
                for row in reader:
                    # Join row cells with delimiter for better context
                    row_text = " | ".join(str(cell).strip() for cell in row if cell)
                    if row_text:
                        text.append(row_text)
            
            return "\n".join(text)
        except Exception as e:
            print(f"   ❌ Error loading CSV: {e}")
            return None


class JSONLoader(DocumentLoader):
    """Loader for JSON files."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text format
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"   ❌ Error loading JSON: {e}")
            return None


class HTMLLoader(DocumentLoader):
    """Loader for HTML files."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except ImportError:
            print(f"   ⚠️  beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
            return None
        except Exception as e:
            print(f"   ❌ Error loading HTML: {e}")
            return None


class ExcelLoader(DocumentLoader):
    """Loader for Excel files (.xlsx, .xls)."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Add sheet name as header
                text.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Convert dataframe to string representation
                # Use to_csv for better text representation
                sheet_text = df.to_csv(index=False, sep='|')
                text.append(sheet_text)
            
            return "\n".join(text)
        except ImportError:
            print(f"   ⚠️  pandas and openpyxl not installed. Install with: pip install pandas openpyxl")
            return None
        except Exception as e:
            print(f"   ❌ Error loading Excel: {e}")
            return None


class MarkdownLoader(DocumentLoader):
    """Loader for Markdown files (.md)."""
    
    @staticmethod
    def load(file_path: Path) -> Optional[str]:
        # Markdown is essentially text, so we can use TextLoader
        return TextLoader.load(file_path)


# Loader registry mapping file extensions to loader classes
LOADER_REGISTRY: Dict[str, type] = {
    '.txt': TextLoader,
    '.md': MarkdownLoader,
    '.pdf': PDFLoader,
    '.docx': DocxLoader,
    '.csv': CSVLoader,
    '.json': JSONLoader,
    '.html': HTMLLoader,
    '.htm': HTMLLoader,
    '.xlsx': ExcelLoader,
    '.xls': ExcelLoader,
}


def get_loader(file_path: Path) -> Optional[type]:
    """Get the appropriate loader for a file based on its extension."""
    extension = file_path.suffix.lower()
    return LOADER_REGISTRY.get(extension)


def load_document(file_path: Path) -> Optional[str]:
    """
    Load a document using the appropriate loader.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Extracted text content or None if loading failed
    """
    loader_class = get_loader(file_path)
    
    if loader_class is None:
        print(f"   ⚠️  No loader available for {file_path.suffix} files")
        return None
    
    return loader_class.load(file_path)


def get_supported_extensions() -> list:
    """Get list of supported file extensions."""
    return list(LOADER_REGISTRY.keys())
